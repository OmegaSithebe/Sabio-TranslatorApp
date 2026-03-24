"""
utils/docx_reader.py

In-place DOCX translation — preserves ALL formatting, images, logos,
table structure, headers, and footers.  Only visible text changes.

Speed approach
--------------
Pass 1: Parse all XML parts, collect every paragraph's text into one list.
Pass 2: Hand the whole list to translate_many() — one batched+parallel call.
Pass 3: Write translations back into the XML nodes and repack the ZIP.
"""

import io
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Callable, Optional
import logging

# ── Register all OOXML namespaces so ET keeps original prefixes ────────────
for _pfx, _uri in {
    "wpc":    "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
    "m":      "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "o":      "urn:schemas-microsoft-com:office:office",
    "r":      "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v":      "urn:schemas-microsoft-com:vml",
    "w":      "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w10":    "urn:schemas-microsoft-com:office:word",
    "w14":    "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15":    "http://schemas.microsoft.com/office/word/2012/wordml",
    "w16":    "http://schemas.microsoft.com/office/word/2018/wordml",
    "w16cex": "http://schemas.microsoft.com/office/word/2018/wordml/cex",
    "w16se":  "http://schemas.microsoft.com/office/word/2015/wordml/symex",
    "wne":    "http://schemas.microsoft.com/office/word/2006/wordml",
    "wpg":    "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    "wpi":    "http://schemas.microsoft.com/office/word/2010/wordprocessingInk",
    "wps":    "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "wp":     "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a":      "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic":    "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "mc":     "http://schemas.openxmlformats.org/markup-compatibility/2006",
}.items():
    try:
        ET.register_namespace(_pfx, _uri)
    except Exception:
        pass

_W      = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WT     = f"{{{_W}}}t"
_WP     = f"{{{_W}}}p"
_WR     = f"{{{_W}}}r"
_WDEL   = f"{{{_W}}}del"
_XML_SP = "{http://www.w3.org/XML/1998/namespace}space"


# ══════════════════════════════════════════════════════════════════════════
# Public API
# ══════════════════════════════════════════════════════════════════════════

def extract_docx_text(file) -> Optional[str]:
    """Return flat text for language detection (structure not preserved)."""
    try:
        from docx import Document
        file.seek(0)
        doc   = Document(file)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as exc:
        logging.error(f"Could not read Word document: {exc}")
        return None


def translate_docx_inplace(
    file,
    translate_fn:      Callable[[str], Optional[str]],
    translate_many_fn: Optional[Callable[[list[str]], dict[str, str]]] = None,
) -> Optional[io.BytesIO]:
    """
    Translate only the text in a DOCX; all formatting is preserved.

    Parameters
    ----------
    file              : Seekable uploaded file.
    translate_fn      : Single-string fallback.
    translate_many_fn : Fast path — takes list[str], returns dict.
    """
    try:
        file.seek(0)
        raw = file.read()
    except Exception as exc:
        raise RuntimeError(f"Could not read file: {exc}") from exc

    try:
        src = zipfile.ZipFile(io.BytesIO(raw), "r")
    except Exception as exc:
        raise ValueError(f"File does not appear to be a valid DOCX: {exc}") from exc

    text_parts = {"word/document.xml"}
    for name in src.namelist():
        if re.match(r"word/(header|footer)\d*\.xml$", name):
            text_parts.add(name)

    # ── Pass 1: parse XML parts, collect all paragraph texts ──────────────
    parsed:       dict[str, ET.Element] = {}
    declarations: dict[str, str]        = {}
    all_texts:    list[str]             = []

    for part in text_parts:
        if part not in src.namelist():
            continue
        raw_xml          = src.read(part).decode("utf-8")
        decl, body       = _split_decl(raw_xml)
        root             = ET.fromstring(body)
        parsed[part]     = root
        declarations[part] = decl
        for para in root.iter(_WP):
            t = _para_text(para)
            if t:
                all_texts.append(t)

    # ── Pass 2: batch-translate all unique strings ─────────────────────────
    unique = list(dict.fromkeys(all_texts))
    if translate_many_fn and unique:
        translations = translate_many_fn(unique)
    else:
        translations = {t: (translate_fn(t) or t) for t in unique}

    # ── Pass 3: apply translations and repack ─────────────────────────────
    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            if item.filename in parsed:
                root = parsed[item.filename]
                for para in root.iter(_WP):
                    _apply_para(para, translations)
                xml_out = (
                    declarations[item.filename]
                    + ET.tostring(root, encoding="unicode", xml_declaration=False)
                )
                dst.writestr(item, xml_out.encode("utf-8"))
            else:
                dst.writestr(item, src.read(item.filename))

    src.close()
    out_buf.seek(0)
    return out_buf


# ══════════════════════════════════════════════════════════════════════════
# Internal XML helpers
# ══════════════════════════════════════════════════════════════════════════

def _para_text(para: ET.Element) -> str:
    """Return the joined visible text of a paragraph (excluding deleted runs)."""
    parts = []
    for run in para.iter(_WR):
        if _inside_del(para, run):
            continue
        for child in run:
            if child.tag == _WT:
                parts.append(child.text or "")
    text = "".join(parts).strip()
    return text if text and not _skip(text) else ""


def _apply_para(para: ET.Element, translations: dict[str, str]) -> None:
    """Write the translated text back into the paragraph's <w:t> nodes."""
    t_nodes = []
    for run in para.iter(_WR):
        if _inside_del(para, run):
            continue
        for child in run:
            if child.tag == _WT:
                t_nodes.append(child)

    if not t_nodes:
        return

    original = "".join(n.text or "" for n in t_nodes).strip()
    if not original:
        return

    translated = translations.get(original, "")
    if not translated or translated.strip() == original:
        return

    t_nodes[0].text = translated
    if translated[:1] == " " or translated[-1:] == " ":
        t_nodes[0].set(_XML_SP, "preserve")
    else:
        t_nodes[0].attrib.pop(_XML_SP, None)

    for node in t_nodes[1:]:
        node.text = ""
        node.attrib.pop(_XML_SP, None)


def _inside_del(root: ET.Element, target: ET.Element) -> bool:
    stack = [(root, False)]
    while stack:
        node, in_del = stack.pop()
        if node is target:
            return in_del
        is_del = node.tag == _WDEL
        for child in node:
            stack.append((child, in_del or is_del))
    return False


def _split_decl(xml: str):
    if xml.startswith("<?xml"):
        idx = xml.index("?>") + 2
        return xml[:idx], xml[idx:]
    return "", xml


def _skip(text: str) -> bool:
    if len(text) <= 1:
        return True
    if re.fullmatch(r"[\d\s.,\-/\\:;()+%$€£@#&*=<>_|~`'\"!?]+", text):
        return True
    if re.match(r"https?://\S+", text):
        return True
    return False


# ── Fallback: plain DOCX builder (Quick Text download only) ───────────────

def create_translated_docx(text: str) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt
    buf = io.BytesIO()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for line in text.split("\n"):
        doc.add_paragraph(line).paragraph_format.space_after = Pt(4)
    doc.save(buf)
    buf.seek(0)
    return buf
