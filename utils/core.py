"""
utils/file_handlers.py

Consolidated file readers for PDF, DOCX, and Excel documents.
All extraction and in-place translation helpers are now in one module.
"""

import io
import re
import zipfile
import xml.etree.ElementTree as ET
import logging
from typing import Callable, Optional

# DOCX namespaces -----------------------------------------------------------
for _pfx, _uri in {
    "wpc":    "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
    "m":      "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "o":      "urn:schemas-microsoft-com:office:office",
    "r":      "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v":      "urn:schemas-microsoft-com:vml",
    "w":      "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w10":    "urn:schemas-microsoft-com:office:word",
    "w14":    "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15":    "http://schemas.microsoft.com/office/word/2012/wordml",
    "w16":    "http://schemas.microsoft.com/office/word/2018/wordml",
    "w16cex": "http://schemas.microsoft.com/office/word/2018/wordml/cex",
    "w16se":  "http://schemas.microsoft.com/office/word/2015/wordml/symex",
    "wne":    "http://schemas.microsoft.com/office/word/2006/wordml",
    "wpg":    "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    "wpi":    "http://schemas.microsoft.com/office/word/2010/wordprocessingInk",
    "wps":    "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "wp":     "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a":      "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic":    "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "mc":     "http://schemas.openxmlformats.org/markup-compatibility/2006",
}.items():
    try:
        ET.register_namespace(_pfx, _uri)
    except Exception:
        pass

_W      = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WT     = f"{{{_W}}}t"
_WP     = f"{{{_W}}}p"
_WR     = f"{{{_W}}}r"
_WDEL   = f"{{{_W}}}del"
_XML_SP = "{http://www.w3.org/XML/1998/namespace}space"

# Excel namespaces ---------------------------------------------------------
for _pfx, _uri in {
    "":       "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
    "r":      "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "mc":     "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "x14ac":  "http://schemas.microsoft.com/office/spreadsheetml/2009/9/ac",
    "xr":     "http://schemas.microsoft.com/office/spreadsheetml/2014/revision",
}.items():
    try:
        ET.register_namespace(_pfx, _uri)
    except Exception:
        pass

_SS   = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
_SI   = f"{{{_SS}}}si"
_T    = f"{{{_SS}}}t"
_RPH  = f"{{{_SS}}}rPh"
_IS   = f"{{{_SS}}}is"
_XSPACE = "{http://www.w3.org/XML/1998/namespace}space"


# File utils constants and functions ---------------------------------------

import os
from typing import Tuple

MAX_FILE_SIZE       = 200 * 1024 * 1024   # 200 MB
ALLOWED_EXTENSIONS  = {".pdf", ".docx", ".xlsx", ".xls"}

_TYPE_LABELS = {
    ".pdf":  "PDF Document",
    ".docx": "Word Document",
    ".xlsx": "Excel Spreadsheet",
    ".xls":  "Excel Spreadsheet (Legacy)",
}
_TYPE_BADGES = {
    ".pdf":  "PDF",
    ".docx": "DOCX",
    ".xlsx": "XLSX",
    ".xls":  "XLS",
}


def get_file_extension(filename: str) -> Optional[str]:
    if not filename:
        return None
    return os.path.splitext(filename)[1].lower()


def allowed_file_type(filename: str) -> bool:
    ext = get_file_extension(filename)
    return ext in ALLOWED_EXTENSIONS if ext else False


def get_file_type_display(extension: str) -> str:
    return _TYPE_LABELS.get(extension.lower(), "Unknown File Type")


def get_file_icon(filename: str) -> str:
    ext = get_file_extension(filename)
    return _TYPE_BADGES.get(ext, "FILE")


def validate_file(file) -> Tuple[bool, str]:
    if file is None:
        return False, "No file provided."
    if hasattr(file, "size") and file.size > MAX_FILE_SIZE:
        mb = file.size / (1024 * 1024)
        return False, f"File exceeds the 200 MB limit ({mb:.1f} MB uploaded)."
    if not allowed_file_type(file.name):
        ext = get_file_extension(file.name) or "unknown"
        return False, f"'{ext}' is not supported. Please upload a PDF, DOCX, or Excel file."
    return True, "Valid"


def format_file_size(size_bytes: int) -> str:
    if size_bytes < 1024:
        return f"{size_bytes} B"
    if size_bytes < 1024 ** 2:
        return f"{size_bytes / 1024:.1f} KB"
    return f"{size_bytes / (1024 ** 2):.1f} MB"


# Translation functions ----------------------------------------------------

import time
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed

SUPPORTED_LANGUAGES: dict[str, str] = {
    "en": "English",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
    "it": "Italian",
    "pt": "Portuguese",
    "nl": "Dutch",
    "ru": "Russian",
    "zh": "Chinese (Simplified)",
    "ja": "Japanese",
    "ko": "Korean",
    "ar": "Arabic",
    "hi": "Hindi",
    "tr": "Turkish",
    "pl": "Polish",
    "uk": "Ukrainian",
    "vi": "Vietnamese",
    "th": "Thai",
    "cs": "Czech",
    "el": "Greek",
    "sv": "Swedish",
    "da": "Danish",
    "fi": "Finnish",
    "no": "Norwegian",
    "hu": "Hungarian",
    "ro": "Romanian",
    "bg": "Bulgarian",
    "id": "Indonesian",
    "ms": "Malay",
}

# ── Tuning knobs ───────────────────────────────────────────────────────────
_BATCH_CHARS   = 4_800
_WORKERS       = 16
_SEP           = " ⏎⏎ "
_SEP_PATTERN   = re.compile(r"\s*⏎⏎\s*")
_RETRIES       = 2
_RETRY_DELAY   = 0.5


def detect_language(text: str) -> Optional[str]:
    if not text or len(text.strip()) < 10:
        return None
    try:
        from langdetect import detect, DetectorFactory
        DetectorFactory.seed = 42
        sample = " ".join(text.split()[:200])
        code   = detect(sample)
        if code.startswith("zh"):
            return "zh"
        code = code.split("-")[0]
        return code if code in SUPPORTED_LANGUAGES else None
    except Exception:
        return None


def translate_text(text: str, source: str, target: str) -> Optional[str]:
    """Translate a single plain-text string (used by the Quick Text panel)."""
    if not text or not text.strip():
        return ""
    source = _normalise(source, text)
    target = target.split("-")[0].lower()
    if source == target:
        return text
    return _translate_one(text, source, target)


def translate_many(
    strings: list[str],
    source: str,
    target: str,
) -> dict[str, str]:
    """
    Translate a list of strings efficiently.

    Returns a dict mapping each original string to its translation.
    Strings that should not be translated are returned unchanged.
    Identical strings are translated only once.

    This is the function called by the PDF / DOCX / XLSX readers.
    """
    if not strings:
        return {}

    source = _normalise(source, " ".join(strings[:20]))
    target = target.split("-")[0].lower()

    # Separate strings that need translation from those that don't
    to_translate: list[str] = []
    result: dict[str, str]  = {}

    for s in strings:
        if not s or not s.strip() or source == target:
            result[s] = s
        else:
            to_translate.append(s)

    if not to_translate:
        return result

    # Deduplicate — translate each unique string only once
    unique = list(dict.fromkeys(to_translate))   # preserves insertion order

    # Pack unique strings into batches
    batches = _make_batches(unique)

    # Translate batches in parallel
    translated_unique: dict[str, str] = {}
    lock = threading.Lock()

    with ThreadPoolExecutor(max_workers=min(_WORKERS, len(batches))) as pool:
        futures = {
            pool.submit(_translate_batch, batch, source, target): batch
            for batch in batches
        }
        for future in as_completed(futures):
            batch_result = future.result()   # dict[original → translated]
            with lock:
                translated_unique.update(batch_result)

    # Map every original string (including duplicates) to its translation
    for s in to_translate:
        result[s] = translated_unique.get(s, s)

    return result


def get_language_name(code: str) -> str:
    return SUPPORTED_LANGUAGES.get(code, code)


def _make_batches(strings: list[str]) -> list[list[str]]:
    """
    Group strings into batches where each batch's joined length stays
    under _BATCH_CHARS.  Single strings longer than _BATCH_CHARS get
    their own batch.
    """
    batches: list[list[str]] = []
    current: list[str]       = []
    current_len              = 0
    sep_len                  = len(_SEP)

    for s in strings:
        s_len = len(s)
        # Would adding this string exceed the batch limit?
        needed = s_len + (sep_len if current else 0)
        if current and current_len + needed > _BATCH_CHARS:
            batches.append(current)
            current     = []
            current_len = 0
        current.append(s)
        current_len += s_len + (sep_len if len(current) > 1 else 0)

    if current:
        batches.append(current)

    return batches


def _translate_batch(batch: list[str], source: str, target: str) -> dict[str, str]:
    """
    Translate a batch of strings using the separator trick.

    We join them with _SEP, send as one request, split on the returned
    separator.  If the split count doesn't match (rare — happens when
    Google occasionally drops the separator), we fall back to translating
    each string individually.
    """
    if len(batch) == 1:
        original    = batch[0]
        translated  = _translate_one(original, source, target)
        return {original: translated or original}

    joined     = _SEP.join(batch)
    translated = _translate_one(joined, source, target)

    if not translated:
        # Full batch failed — fall back one-by-one
        return {s: (_translate_one(s, source, target) or s) for s in batch}

    parts = _SEP_PATTERN.split(translated)

    if len(parts) == len(batch):
        return {orig: trans for orig, trans in zip(batch, parts)}

    # Separator count mismatch — fall back one-by-one to stay correct
    result = {}
    for s in batch:
        result[s] = _translate_one(s, source, target) or s
    return result


def _translate_one(text: str, source: str, target: str) -> Optional[str]:
    """Single-string translation with retry logic."""
    try:
        from deep_translator import GoogleTranslator
    except ImportError:
        raise ImportError("deep-translator not installed. Run: pip install deep-translator")

    for attempt in range(_RETRIES):
        try:
            result = GoogleTranslator(source=source, target=target).translate(text)
            if result:
                return result
        except Exception as exc:
            err = str(exc)
            if "429" in err:
                time.sleep(_RETRY_DELAY * (attempt + 2))   # back off harder on rate limit
            elif attempt < _RETRIES - 1:
                time.sleep(_RETRY_DELAY)
            else:
                return None
    return None


def _normalise(source: str, sample_text: str) -> str:
    """Resolve 'auto' to a real language code."""
    if source == "auto":
        return detect_language(sample_text) or "en"
    return source.split("-")[0].lower()


# Public API ----------------------------------------------------------------

def extract_docx_text(file) -> Optional[str]:
    """Return flat text for language detection (structure not preserved)."""
    try:
        from docx import Document
        file.seek(0)
        doc = Document(file)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as exc:
        logging.error(f"Could not read Word document: {exc}")
        return None


def translate_docx_inplace(
    file,
    translate_fn: Callable[[str], Optional[str]],
    translate_many_fn: Optional[Callable[[list[str]], dict[str, str]]] = None,
) -> Optional[io.BytesIO]:
    """Translate only the text in a DOCX; all formatting is preserved."""
    try:
        file.seek(0)
        raw = file.read()
    except Exception as exc:
        raise RuntimeError(f"Could not read file: {exc}") from exc

    try:
        src = zipfile.ZipFile(io.BytesIO(raw), "r")
    except Exception as exc:
        raise ValueError(f"File does not appear to be a valid DOCX: {exc}") from exc

    text_parts = {"word/document.xml"}
    for name in src.namelist():
        if re.match(r"word/(header|footer)\d*\.xml$", name):
            text_parts.add(name)

    parsed: dict[str, ET.Element] = {}
    declarations: dict[str, str] = {}
    all_texts: list[str] = []

    for part in text_parts:
        if part not in src.namelist():
            continue
        raw_xml = src.read(part).decode("utf-8")
        decl, body = _split_decl(raw_xml)
        root = ET.fromstring(body)
        parsed[part] = root
        declarations[part] = decl
        for para in root.iter(_WP):
            t = _para_text(para)
            if t:
                all_texts.append(t)

    unique = list(dict.fromkeys(all_texts))
    if translate_many_fn and unique:
        translations = translate_many_fn(unique)
    else:
        translations = {t: (translate_fn(t) or t) for t in unique}

    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            if item.filename in parsed:
                root = parsed[item.filename]
                for para in root.iter(_WP):
                    _apply_para(para, translations)
                xml_out = declarations[item.filename] + ET.tostring(root, encoding="unicode", xml_declaration=False)
                dst.writestr(item, xml_out.encode("utf-8"))
            else:
                dst.writestr(item, src.read(item.filename))

    src.close()
    out_buf.seek(0)
    return out_buf


def create_translated_docx(text: str) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt
    buf = io.BytesIO()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for line in text.split("\n"):
        doc.add_paragraph(line).paragraph_format.space_after = Pt(4)
    doc.save(buf)
    buf.seek(0)
    return buf


def extract_excel_text(file) -> Optional[str]:
    """Return flat text for language detection only."""
    try:
        import pandas as pd
        file.seek(0)
        xf = pd.ExcelFile(file)
        parts = []
        for sheet in xf.sheet_names:
            file.seek(0)
            df = pd.read_excel(file, sheet_name=sheet, dtype=str).fillna("")
            for col in df.columns:
                if str(col).strip():
                    parts.append(str(col))
            for _, row in df.iterrows():
                for val in row:
                    if str(val).strip():
                        parts.append(str(val))
        return "\n".join(parts)
    except Exception as exc:
        logging.error(f"Could not read Excel file: {exc}")
        return None


def translate_xlsx_inplace(
    file,
    translate_fn: Callable[[str], Optional[str]],
    translate_many_fn: Optional[Callable[[list[str]], dict[str, str]]] = None,
) -> Optional[io.BytesIO]:
    """Translate only the text in an XLSX; all formatting is preserved."""
    try:
        file.seek(0)
        raw = file.read()
    except Exception as exc:
        raise RuntimeError(f"Could not read file: {exc}") from exc

    if raw[:2] == b"\xd0\xcf":
        raise ValueError(
            "Legacy .xls files cannot be translated in-place (binary format). "
            "Open in Excel, save as .xlsx, then re-upload."
        )

    try:
        src = zipfile.ZipFile(io.BytesIO(raw), "r")
    except Exception as exc:
        raise ValueError(f"File does not appear to be a valid XLSX: {exc}") from exc

    ss_data: Optional[bytes] = None
    ss_root: Optional[ET.Element] = None
    ss_decl: str = ""
    all_texts: list[str] = []

    if "xl/sharedStrings.xml" in src.namelist():
        ss_data = src.read("xl/sharedStrings.xml")
        ss_xml = ss_data.decode("utf-8")
        ss_decl, ss_body = _split_decl(ss_xml)
        ss_root = ET.fromstring(ss_body)
        for si in ss_root.iter(_SI):
            t = _si_text(si)
            if t:
                all_texts.append(t)

    sheet_parts: dict[str, tuple[ET.Element, str, bytes]] = {}
    for name in src.namelist():
        if re.match(r"xl/worksheets/sheet\d+\.xml$", name):
            raw_xml = src.read(name)
            xml_str = raw_xml.decode("utf-8")
            decl, body = _split_decl(xml_str)
            root = ET.fromstring(body)
            sheet_parts[name] = (root, decl, raw_xml)
            for is_node in root.iter(_IS):
                t = _node_text(is_node)
                if t:
                    all_texts.append(t)

    unique = list(dict.fromkeys(all_texts))
    if translate_many_fn and unique:
        translations = translate_many_fn(unique)
    else:
        translations = {t: (translate_fn(t) or t) for t in unique}

    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)

            if item.filename == "xl/sharedStrings.xml" and ss_root is not None:
                for si in ss_root.iter(_SI):
                    _apply_si(si, translations)
                xml_out = ss_decl + ET.tostring(ss_root, encoding="unicode", xml_declaration=False)
                data = xml_out.encode("utf-8")

            elif item.filename in sheet_parts:
                root, decl, _ = sheet_parts[item.filename]
                for is_node in root.iter(_IS):
                    _apply_inline(is_node, translations)
                xml_out = decl + ET.tostring(root, encoding="unicode", xml_declaration=False)
                data = xml_out.encode("utf-8")

            dst.writestr(item, data)

    src.close()
    out_buf.seek(0)
    return out_buf


def extract_pdf_text(file) -> Optional[str]:
    """Extract plain text for language detection."""
    try:
        import fitz
        file.seek(0)
        doc = fitz.open(stream=file.read(), filetype="pdf")
        parts = [page.get_text().strip() for page in doc if page.get_text().strip()]
        doc.close()
        return "\n\n".join(parts) or ""
    except Exception as exc:
        logging.error(f"Could not read PDF: {exc}")
        return None


def translate_pdf_inplace(
    file,
    translate_fn: Callable[[str], Optional[str]],
    translate_many_fn: Optional[Callable[[list[str]], dict[str, str]]] = None,
) -> Optional[io.BytesIO]:
    """Translate a PDF in-place; preserves layout and graphics."""
    try:
        import fitz
    except ImportError:
        raise ImportError("PyMuPDF not installed. Run: pip install pymupdf")

    try:
        file.seek(0)
        doc = fitz.open(stream=file.read(), filetype="pdf")
    except Exception as exc:
        raise RuntimeError(f"Could not open PDF: {exc}") from exc

    has_text = any(page.get_text().strip() for page in doc)

    if not has_text:
        logging.warning(
            "This PDF has no selectable text (fully scanned). "
            "Install pytesseract and Pillow for OCR support."
        )
        result = _translate_scanned(doc, translate_fn)
        doc.close()
        return result

    try:
        page_spans: list[list[dict]] = []
        all_originals: list[str] = []

        for page in doc:
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
            spans: list[dict] = []

            for block in blocks:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span["text"].strip()
                        if not text or _skip(text):
                            continue
                        spans.append({
                            "original": text,
                            "rect": fitz.Rect(span["bbox"]),
                            "font_size": span["size"],
                            "color": _unpack_color(span.get("color", 0)),
                            "font_name": _safe_font(span.get("font", "helv")),
                        })
                        all_originals.append(text)

            page_spans.append(spans)

        if not all_originals:
            buf = io.BytesIO()
            doc.save(buf, garbage=2, deflate=True)
            doc.close()
            buf.seek(0)
            return buf

        unique = list(dict.fromkeys(all_originals))

        if translate_many_fn:
            translations = translate_many_fn(unique)
        else:
            translations = {t: (translate_fn(t) or t) for t in unique}

        for page, spans in zip(doc, page_spans):
            to_apply = []
            for span in spans:
                translated = translations.get(span["original"], "").strip()
                if not translated or translated == span["original"]:
                    continue
                to_apply.append({**span, "translated": translated})

            for item in to_apply:
                page.add_redact_annot(item["rect"], fill=(1, 1, 1))
            page.apply_redactions()

            for item in to_apply:
                _insert_text(page, item)

        buf = io.BytesIO()
        doc.save(buf, garbage=2, deflate=True)
        doc.close()
        buf.seek(0)
        return buf

    except Exception as exc:
        doc.close()
        raise RuntimeError(f"Error translating PDF: {exc}") from exc


def _insert_text(page, item: dict) -> None:
    kwargs = dict(
        fontsize=item["font_size"],
        color=item["color"],
        fontname=item["font_name"],
    )
    try:
        page.insert_text(item["rect"].tl, item["translated"], **kwargs)
    except Exception:
        try:
            page.insert_text(item["rect"].tl, item["translated"],
                             fontsize=item["font_size"], color=item["color"],
                             fontname="helv")
        except Exception:
            pass


def _translate_scanned(doc, translate_fn: Callable[[str], Optional[str]]) -> Optional[io.BytesIO]:
    try:
        import pytesseract
        from PIL import Image
        import io as _io
    except ImportError:
        raise ImportError(
            "Scanned PDF detected. Install pytesseract and Pillow: "
            "pip install pytesseract Pillow"
        )

    for page in doc:
        mat = page.get_pixmap(dpi=200)
        img = Image.open(_io.BytesIO(mat.tobytes("png")))
        ocr_text = pytesseract.image_to_string(img).strip()
        if not ocr_text:
            continue
        translated = translate_fn(ocr_text)
        if translated:
            page.insert_textbox(page.rect, translated,
                                fontsize=11, fontname="helv", color=(0, 0, 0))

    buf = io.BytesIO()
    doc.save(buf, garbage=2, deflate=True)
    buf.seek(0)
    return buf


def is_excel_file(filename: str) -> bool:
    return filename.lower().endswith((".xlsx", ".xls", ".xlsm"))


def _para_text(para: ET.Element) -> str:
    parts = []
    for run in para.iter(_WR):
        if _inside_del(para, run):
            continue
        for child in run:
            if child.tag == _WT:
                parts.append(child.text or "")
    text = "".join(parts).strip()
    return text if text and not _skip(text) else ""


def _apply_para(para: ET.Element, translations: dict[str, str]) -> None:
    t_nodes = []
    for run in para.iter(_WR):
        if _inside_del(para, run):
            continue
        for child in run:
            if child.tag == _WT:
                t_nodes.append(child)

    if not t_nodes:
        return

    original = "".join(n.text or "" for n in t_nodes).strip()
    if not original:
        return

    translated = translations.get(original, "")
    if not translated or translated.strip() == original:
        return

    t_nodes[0].text = translated
    if translated[:1] == " " or translated[-1:] == " ":
        t_nodes[0].set(_XML_SP, "preserve")
    else:
        t_nodes[0].attrib.pop(_XML_SP, None)

    for node in t_nodes[1:]:
        node.text = ""
        node.attrib.pop(_XML_SP, None)


def _inside_del(root: ET.Element, target: ET.Element) -> bool:
    stack = [(root, False)]
    while stack:
        node, in_del = stack.pop()
        if node is target:
            return in_del
        is_del = node.tag == _WDEL
        for child in node:
            stack.append((child, in_del or is_del))
    return False


def _si_text(si: ET.Element) -> str:
    parts = []
    for t in si.iter(_T):
        if not _inside_rph(si, t):
            parts.append(t.text or "")
    text = "".join(parts).strip()
    return text if text and not _skip(text) else ""


def _node_text(node: ET.Element) -> str:
    parts = [t.text or "" for t in node.iter(_T)]
    text = "".join(parts).strip()
    return text if text and not _skip(text) else ""


def _apply_si(si: ET.Element, translations: dict[str, str]) -> None:
    t_nodes = [t for t in si.iter(_T) if not _inside_rph(si, t)]
    if not t_nodes:
        return
    original = "".join(t.text or "" for t in t_nodes).strip()
    translated = translations.get(original, "")
    if not translated or translated.strip() == original:
        return
    t_nodes[0].text = translated
    _set_space(t_nodes[0], translated)
    for t in t_nodes[1:]:
        t.text = ""
        t.attrib.pop(_XSPACE, None)


def _apply_inline(is_node: ET.Element, translations: dict[str, str]) -> None:
    t_nodes = list(is_node.iter(_T))
    if not t_nodes:
        return
    original = "".join(t.text or "" for t in t_nodes).strip()
    translated = translations.get(original, "")
    if not translated or translated.strip() == original:
        return
    t_nodes[0].text = translated
    _set_space(t_nodes[0], translated)
    for t in t_nodes[1:]:
        t.text = ""


def _inside_rph(root: ET.Element, target: ET.Element) -> bool:
    stack = [(root, False)]
    while stack:
        node, in_rph = stack.pop()
        if node is target:
            return in_rph
        for child in node:
            stack.append((child, in_rph or node.tag == _RPH))
    return False


def _set_space(node: ET.Element, text: str) -> None:
    if text and (text[0] == " " or text[-1] == " "):
        node.set(_XSPACE, "preserve")
    else:
        node.attrib.pop(_XSPACE, None)


def _split_decl(xml: str):
    if xml.startswith("<?xml"):
        idx = xml.index("?>") + 2
        return xml[:idx], xml[idx:]
    return "", xml


def _skip(text: str) -> bool:
    if len(text) <= 1:
        return True
    if re.fullmatch(r"[\d\s.,\-/\\:;()+%$€£@#&*=<>_|~`'\"!?]+", text):
        return True
    if re.match(r"https?://\S+", text):
        return True
    return False


def _unpack_color(packed: int) -> tuple:
    r = ((packed >> 16) & 0xFF) / 255.0
    g = ((packed >>  8) & 0xFF) / 255.0
    b = ((packed      ) & 0xFF) / 255.0
    return (r, g, b)


def _safe_font(font_name: str) -> str:
    name = font_name.lower()
    if any(k in name for k in ("courier", "cour", "mono", "consol")):
        if "bold" in name and "italic" in name: return "courbi"
        if "bold" in name:                       return "courb"
        if "italic" in name or "oblique" in name: return "couri"
        return "cour"
    if any(k in name for k in ("times", "serif", "georgia")):
        if "bold" in name and "italic" in name: return "timbi"
        if "bold" in name:                       return "timb"
        if "italic" in name or "oblique" in name: return "timi"
        return "timr"
    if any(k in name for k in ("helvetica", "arial", "sans")):
        if "bold" in name and "italic" in name: return "heboit"
        if "bold" in name:                       return "hebo"
        if "italic" in name or "oblique" in name: return "hebi"
        return "helv"
    return "helv"

# DOCX namespaces -----------------------------------------------------------
for _pfx, _uri in {
    "wpc":    "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
    "m":      "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "o":      "urn:schemas-microsoft-com:office:office",
    "r":      "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v":      "urn:schemas-microsoft-com:vml",
    "w":      "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w10":    "urn:schemas-microsoft-com:office:word",
    "w14":    "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15":    "http://schemas.microsoft.com/office/word/2012/wordml",
    "w16":    "http://schemas.microsoft.com/office/word/2018/wordml",
    "w16cex": "http://schemas.microsoft.com/office/word/2018/wordml/cex",
    "w16se":  "http://schemas.microsoft.com/office/word/2015/wordml/symex",
    "wne":    "http://schemas.microsoft.com/office/word/2006/wordml",
    "wpg":    "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    "wpi":    "http://schemas.microsoft.com/office/word/2010/wordprocessingInk",
    "wps":    "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "wp":     "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a":      "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic":    "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "mc":     "http://schemas.openxmlformats.org/markup-compatibility/2006",
}.items():
    try:
        ET.register_namespace(_pfx, _uri)
    except Exception:
        pass

_W      = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WT     = f"{{{_W}}}t"
_WP     = f"{{{_W}}}p"
_WR     = f"{{{_W}}}r"
_WDEL   = f"{{{_W}}}del"
_XML_SP = "{http://www.w3.org/XML/1998/namespace}space"

# Excel namespaces ---------------------------------------------------------
for _pfx, _uri in {
    "":       "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
    "r":      "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "mc":     "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "x14ac":  "http://schemas.microsoft.com/office/spreadsheetml/2009/9/ac",
    "xr":     "http://schemas.microsoft.com/office/spreadsheetml/2014/revision",
}.items():
    try:
        ET.register_namespace(_pfx, _uri)
    except Exception:
        pass

_SS   = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
_SI   = f"{{{_SS}}}si"
_T    = f"{{{_SS}}}t"
_RPH  = f"{{{_SS}}}rPh"
_IS   = f"{{{_SS}}}is"
_XSPACE = "{http://www.w3.org/XML/1998/namespace}space"


# Public API ----------------------------------------------------------------

def extract_docx_text(file) -> Optional[str]:
    """Return flat text for language detection (structure not preserved)."""
    try:
        from docx import Document
        file.seek(0)
        doc = Document(file)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as exc:
        logging.error(f"Could not read Word document: {exc}")
        return None


def translate_docx_inplace(
    file,
    translate_fn: Callable[[str], Optional[str]],
    translate_many_fn: Optional[Callable[[list[str]], dict[str, str]]] = None,
) -> Optional[io.BytesIO]:
    """Translate only the text in a DOCX; all formatting is preserved."""
    try:
        file.seek(0)
        raw = file.read()
    except Exception as exc:
        raise RuntimeError(f"Could not read file: {exc}") from exc

    try:
        src = zipfile.ZipFile(io.BytesIO(raw), "r")
    except Exception as exc:
        raise ValueError(f"File does not appear to be a valid DOCX: {exc}") from exc

    text_parts = {"word/document.xml"}
    for name in src.namelist():
        if re.match(r"word/(header|footer)\d*\.xml$", name):
            text_parts.add(name)

    parsed: dict[str, ET.Element] = {}
    declarations: dict[str, str] = {}
    all_texts: list[str] = []

    for part in text_parts:
        if part not in src.namelist():
            continue
        raw_xml = src.read(part).decode("utf-8")
        decl, body = _split_decl(raw_xml)
        root = ET.fromstring(body)
        parsed[part] = root
        declarations[part] = decl
        for para in root.iter(_WP):
            t = _para_text(para)
            if t:
                all_texts.append(t)

    unique = list(dict.fromkeys(all_texts))
    if translate_many_fn and unique:
        translations = translate_many_fn(unique)
    else:
        translations = {t: (translate_fn(t) or t) for t in unique}

    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            if item.filename in parsed:
                root = parsed[item.filename]
                for para in root.iter(_WP):
                    _apply_para(para, translations)
                xml_out = declarations[item.filename] + ET.tostring(root, encoding="unicode", xml_declaration=False)
                dst.writestr(item, xml_out.encode("utf-8"))
            else:
                dst.writestr(item, src.read(item.filename))

    src.close()
    out_buf.seek(0)
    return out_buf


def create_translated_docx(text: str) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt
    buf = io.BytesIO()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for line in text.split("\n"):
        doc.add_paragraph(line).paragraph_format.space_after = Pt(4)
    doc.save(buf)
    buf.seek(0)
    return buf


def extract_excel_text(file) -> Optional[str]:
    """Return flat text for language detection only."""
    try:
        import pandas as pd
        file.seek(0)
        xf = pd.ExcelFile(file)
        parts = []
        for sheet in xf.sheet_names:
            file.seek(0)
            df = pd.read_excel(file, sheet_name=sheet, dtype=str).fillna("")
            for col in df.columns:
                if str(col).strip():
                    parts.append(str(col))
            for _, row in df.iterrows():
                for val in row:
                    if str(val).strip():
                        parts.append(str(val))
        return "\n".join(parts)
    except Exception as exc:
        logging.error(f"Could not read Excel file: {exc}")
        return None


def translate_xlsx_inplace(
    file,
    translate_fn: Callable[[str], Optional[str]],
    translate_many_fn: Optional[Callable[[list[str]], dict[str, str]]] = None,
) -> Optional[io.BytesIO]:
    """Translate only the text in an XLSX; all formatting is preserved."""
    try:
        file.seek(0)
        raw = file.read()
    except Exception as exc:
        raise RuntimeError(f"Could not read file: {exc}") from exc

    if raw[:2] == b"\xd0\xcf":
        raise ValueError(
            "Legacy .xls files cannot be translated in-place (binary format). "
            "Open in Excel, save as .xlsx, then re-upload."
        )

    try:
        src = zipfile.ZipFile(io.BytesIO(raw), "r")
    except Exception as exc:
        raise ValueError(f"File does not appear to be a valid XLSX: {exc}") from exc

    ss_data: Optional[bytes] = None
    ss_root: Optional[ET.Element] = None
    ss_decl: str = ""
    all_texts: list[str] = []

    if "xl/sharedStrings.xml" in src.namelist():
        ss_data = src.read("xl/sharedStrings.xml")
        ss_xml = ss_data.decode("utf-8")
        ss_decl, ss_body = _split_decl(ss_xml)
        ss_root = ET.fromstring(ss_body)
        for si in ss_root.iter(_SI):
            t = _si_text(si)
            if t:
                all_texts.append(t)

    sheet_parts: dict[str, tuple[ET.Element, str, bytes]] = {}
    for name in src.namelist():
        if re.match(r"xl/worksheets/sheet\d+\.xml$", name):
            raw_xml = src.read(name)
            xml_str = raw_xml.decode("utf-8")
            decl, body = _split_decl(xml_str)
            root = ET.fromstring(body)
            sheet_parts[name] = (root, decl, raw_xml)
            for is_node in root.iter(_IS):
                t = _node_text(is_node)
                if t:
                    all_texts.append(t)

    unique = list(dict.fromkeys(all_texts))
    if translate_many_fn and unique:
        translations = translate_many_fn(unique)
    else:
        translations = {t: (translate_fn(t) or t) for t in unique}

    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)

            if item.filename == "xl/sharedStrings.xml" and ss_root is not None:
                for si in ss_root.iter(_SI):
                    _apply_si(si, translations)
                xml_out = ss_decl + ET.tostring(ss_root, encoding="unicode", xml_declaration=False)
                data = xml_out.encode("utf-8")

            elif item.filename in sheet_parts:
                root, decl, _ = sheet_parts[item.filename]
                for is_node in root.iter(_IS):
                    _apply_inline(is_node, translations)
                xml_out = decl + ET.tostring(root, encoding="unicode", xml_declaration=False)
                data = xml_out.encode("utf-8")

            dst.writestr(item, data)

    src.close()
    out_buf.seek(0)
    return out_buf


def extract_pdf_text(file) -> Optional[str]:
    """Extract plain text for language detection."""
    try:
        import fitz
        file.seek(0)
        doc = fitz.open(stream=file.read(), filetype="pdf")
        parts = [page.get_text().strip() for page in doc if page.get_text().strip()]
        doc.close()
        return "\n\n".join(parts) or ""
    except Exception as exc:
        logging.error(f"Could not read PDF: {exc}")
        return None


def translate_pdf_inplace(
    file,
    translate_fn: Callable[[str], Optional[str]],
    translate_many_fn: Optional[Callable[[list[str]], dict[str, str]]] = None,
) -> Optional[io.BytesIO]:
    """Translate a PDF in-place; preserves layout and graphics."""
    try:
        import fitz
    except ImportError:
        raise ImportError("PyMuPDF not installed. Run: pip install pymupdf")

    try:
        file.seek(0)
        doc = fitz.open(stream=file.read(), filetype="pdf")
    except Exception as exc:
        raise RuntimeError(f"Could not open PDF: {exc}") from exc

    has_text = any(page.get_text().strip() for page in doc)

    if not has_text:
        logging.warning(
            "This PDF has no selectable text (fully scanned). "
            "Install pytesseract and Pillow for OCR support."
        )
        result = _translate_scanned(doc, translate_fn)
        doc.close()
        return result

    try:
        page_spans: list[list[dict]] = []
        all_originals: list[str] = []

        for page in doc:
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
            spans: list[dict] = []

            for block in blocks:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span["text"].strip()
                        if not text or _skip(text):
                            continue
                        spans.append({
                            "original": text,
                            "rect": fitz.Rect(span["bbox"]),
                            "font_size": span["size"],
                            "color": _unpack_color(span.get("color", 0)),
                            "font_name": _safe_font(span.get("font", "helv")),
                        })
                        all_originals.append(text)

            page_spans.append(spans)

        if not all_originals:
            buf = io.BytesIO()
            doc.save(buf, garbage=2, deflate=True)
            doc.close()
            buf.seek(0)
            return buf

        unique = list(dict.fromkeys(all_originals))

        if translate_many_fn:
            translations = translate_many_fn(unique)
        else:
            translations = {t: (translate_fn(t) or t) for t in unique}

        for page, spans in zip(doc, page_spans):
            to_apply = []
            for span in spans:
                translated = translations.get(span["original"], "").strip()
                if not translated or translated == span["original"]:
                    continue
                to_apply.append({**span, "translated": translated})

            for item in to_apply:
                page.add_redact_annot(item["rect"], fill=(1, 1, 1))
            page.apply_redactions()

            for item in to_apply:
                _insert_text(page, item)

        buf = io.BytesIO()
        doc.save(buf, garbage=2, deflate=True)
        doc.close()
        buf.seek(0)
        return buf

    except Exception as exc:
        doc.close()
        raise RuntimeError(f"Error translating PDF: {exc}") from exc


def _insert_text(page, item: dict) -> None:
    kwargs = dict(
        fontsize=item["font_size"],
        color=item["color"],
        fontname=item["font_name"],
    )
    try:
        page.insert_text(item["rect"].tl, item["translated"], **kwargs)
    except Exception:
        try:
            page.insert_text(item["rect"].tl, item["translated"],
                             fontsize=item["font_size"], color=item["color"],
                             fontname="helv")
        except Exception:
            pass


def _translate_scanned(doc, translate_fn: Callable[[str], Optional[str]]) -> Optional[io.BytesIO]:
    try:
        import pytesseract
        from PIL import Image
        import io as _io
    except ImportError:
        raise ImportError(
            "Scanned PDF detected. Install pytesseract and Pillow: "
            "pip install pytesseract Pillow"
        )

    for page in doc:
        mat = page.get_pixmap(dpi=200)
        img = Image.open(_io.BytesIO(mat.tobytes("png")))
        ocr_text = pytesseract.image_to_string(img).strip()
        if not ocr_text:
            continue
        translated = translate_fn(ocr_text)
        if translated:
            page.insert_textbox(page.rect, translated,
                                fontsize=11, fontname="helv", color=(0, 0, 0))

    buf = io.BytesIO()
    doc.save(buf, garbage=2, deflate=True)
    buf.seek(0)
    return buf


def is_excel_file(filename: str) -> bool:
    return filename.lower().endswith((".xlsx", ".xls", ".xlsm"))


def _para_text(para: ET.Element) -> str:
    parts = []
    for run in para.iter(_WR):
        if _inside_del(para, run):
            continue
        for child in run:
            if child.tag == _WT:
                parts.append(child.text or "")
    text = "".join(parts).strip()
    return text if text and not _skip(text) else ""


def _apply_para(para: ET.Element, translations: dict[str, str]) -> None:
    t_nodes = []
    for run in para.iter(_WR):
        if _inside_del(para, run):
            continue
        for child in run:
            if child.tag == _WT:
                t_nodes.append(child)

    if not t_nodes:
        return

    original = "".join(n.text or "" for n in t_nodes).strip()
    if not original:
        return

    translated = translations.get(original, "")
    if not translated or translated.strip() == original:
        return

    t_nodes[0].text = translated
    if translated[:1] == " " or translated[-1:] == " ":
        t_nodes[0].set(_XML_SP, "preserve")
    else:
        t_nodes[0].attrib.pop(_XML_SP, None)

    for node in t_nodes[1:]:
        node.text = ""
        node.attrib.pop(_XML_SP, None)


def _inside_del(root: ET.Element, target: ET.Element) -> bool:
    stack = [(root, False)]
    while stack:
        node, in_del = stack.pop()
        if node is target:
            return in_del
        is_del = node.tag == _WDEL
        for child in node:
            stack.append((child, in_del or is_del))
    return False


def _si_text(si: ET.Element) -> str:
    parts = []
    for t in si.iter(_T):
        if not _inside_rph(si, t):
            parts.append(t.text or "")
    text = "".join(parts).strip()
    return text if text and not _skip(text) else ""


def _node_text(node: ET.Element) -> str:
    parts = [t.text or "" for t in node.iter(_T)]
    text = "".join(parts).strip()
    return text if text and not _skip(text) else ""


def _apply_si(si: ET.Element, translations: dict[str, str]) -> None:
    t_nodes = [t for t in si.iter(_T) if not _inside_rph(si, t)]
    if not t_nodes:
        return
    original = "".join(t.text or "" for t in t_nodes).strip()
    translated = translations.get(original, "")
    if not translated or translated.strip() == original:
        return
    t_nodes[0].text = translated
    _set_space(t_nodes[0], translated)
    for t in t_nodes[1:]:
        t.text = ""
        t.attrib.pop(_XSPACE, None)


def _apply_inline(is_node: ET.Element, translations: dict[str, str]) -> None:
    t_nodes = list(is_node.iter(_T))
    if not t_nodes:
        return
    original = "".join(t.text or "" for t in t_nodes).strip()
    translated = translations.get(original, "")
    if not translated or translated.strip() == original:
        return
    t_nodes[0].text = translated
    _set_space(t_nodes[0], translated)
    for t in t_nodes[1:]:
        t.text = ""


def _inside_rph(root: ET.Element, target: ET.Element) -> bool:
    stack = [(root, False)]
    while stack:
        node, in_rph = stack.pop()
        if node is target:
            return in_rph
        for child in node:
            stack.append((child, in_rph or node.tag == _RPH))
    return False


def _set_space(node: ET.Element, text: str) -> None:
    if text and (text[0] == " " or text[-1] == " "):
        node.set(_XSPACE, "preserve")
    else:
        node.attrib.pop(_XSPACE, None)


def _split_decl(xml: str):
    if xml.startswith("<?xml"):
        idx = xml.index("?>") + 2
        return xml[:idx], xml[idx:]
    return "", xml


def _skip(text: str) -> bool:
    if len(text) <= 1:
        return True
    if re.fullmatch(r"[\d\s.,\-/\\:;()+%$€£@#&*=<>_|~`'\"!?]+", text):
        return True
    if re.match(r"https?://\S+", text):
        return True
    return False
