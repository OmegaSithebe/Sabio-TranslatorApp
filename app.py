"""
app.py — Sabio Language Translator
Flask backend serving the HTML/CSS/JS frontend.
All translation logic lives in utils/; this file only handles HTTP routing.
"""

import io
import os
import uuid
import logging
from datetime import datetime
from functools import partial

from flask import Flask, request, jsonify, send_file, render_template

from utils.translator import (
    detect_language,
    translate_text,
    translate_many,
    get_language_name,
)
from utils.file_utils import (
    get_file_extension,
    get_file_icon,
    format_file_size,
    allowed_file_type,
    MAX_FILE_SIZE,
)
from utils.pdf_reader   import extract_pdf_text,   translate_pdf_inplace
from utils.docx_reader  import extract_docx_text,  translate_docx_inplace
from utils.excel_reader import extract_excel_text, translate_xlsx_inplace

app = Flask(__name__)
app.secret_key = os.urandom(24)

# In-memory session store.  Each entry keyed by a UUID string.
# { session_id: { bytes, filename, ext, detected_code,
#                 translated_bytes, target_lang, target_name, preview } }
_store: dict[str, dict] = {}


# ── helpers ────────────────────────────────────────────────────────────────

def _translate_one_fn(src: str, tgt: str):
    def fn(text: str) -> str:
        if not text or not text.strip():
            return text
        return translate_text(text, src, tgt) or text
    return fn


def _translate_many_fn(src: str, tgt: str):
    def fn(strings: list[str]) -> dict[str, str]:
        return translate_many(strings, src, tgt)
    return fn


def _extract_preview(data: bytes | io.BytesIO, ext: str) -> str:
    try:
        raw = data.read() if hasattr(data, "read") else data
        if ext == "pdf":
            import fitz
            doc  = fitz.open(stream=raw, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc[:3])
            doc.close()
            return text[:1200].strip()
        if ext == "docx":
            from docx import Document
            doc  = Document(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs[:30] if p.text.strip())
            return text[:1200].strip()
    except Exception:
        pass
    try:
        return data.decode("utf-8", errors="ignore")[:1200] if isinstance(data, bytes) else ""
    except Exception:
        return ""


def _bytes_to_text(data: bytes, ext: str) -> str:
    try:
        if ext == "pdf":
            import fitz
            doc  = fitz.open(stream=data, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            return text
        if ext == "docx":
            from docx import Document
            doc = Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        pass
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _mime(fmt: str) -> str:
    return {
        "pdf":  "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "txt":  "text/plain; charset=utf-8",
        "html": "text/html; charset=utf-8",
        "rtf":  "application/rtf",
        "odt":  "application/vnd.oasis.opendocument.text",
    }.get(fmt, "application/octet-stream")


# ── routes ─────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "Empty filename"}), 400

    if not allowed_file_type(f.filename):
        ext = get_file_extension(f.filename) or "unknown"
        return jsonify({
            "error": f"'{ext}' is not supported. Please upload PDF, DOCX, or XLSX."
        }), 400

    file_bytes = f.read()

    if len(file_bytes) > MAX_FILE_SIZE:
        mb = len(file_bytes) / (1024 * 1024)
        return jsonify({"error": f"File exceeds 200 MB limit ({mb:.1f} MB)."}), 400

    ext = (get_file_extension(f.filename) or "").lstrip(".")

    # Detect language from extracted text
    bio = io.BytesIO(file_bytes)
    bio.name = f.filename
    sample = ""
    try:
        if ext == "pdf":
            sample = extract_pdf_text(bio) or ""
        elif ext == "docx":
            sample = extract_docx_text(bio) or ""
        elif ext in ("xlsx", "xls"):
            sample = extract_excel_text(bio) or ""
        else:
            sample = file_bytes.decode("utf-8", errors="ignore")[:2000]
    except Exception as exc:
        logging.warning(f"Text extraction failed: {exc}")

    detected_code = detect_language(sample) if sample else None
    confidence = 0
    if sample and detected_code:
        try:
            from langdetect import detect_langs, DetectorFactory
            DetectorFactory.seed = 42
            langs = detect_langs(sample[:500])
            if langs and langs[0].lang.split("-")[0] == detected_code.split("-")[0]:
                confidence = int(langs[0].prob * 100)
        except Exception:
            confidence = 85  # reasonable default

    detected_code = detected_code or "en"
    detected_name = get_language_name(detected_code)

    session_id = str(uuid.uuid4())
    _store[session_id] = {
        "bytes":          file_bytes,
        "filename":       f.filename,
        "ext":            ext,
        "detected_code":  detected_code,
    }

    return jsonify({
        "session_id":    session_id,
        "filename":      f.filename,
        "size":          format_file_size(len(file_bytes)),
        "detected_lang": detected_name,
        "detected_code": detected_code,
        "confidence":    confidence,
        "icon":          get_file_icon(f.filename),
    })


@app.route("/api/translate", methods=["POST"])
def translate():
    data = request.get_json(silent=True) or {}
    session_id  = data.get("session_id")
    target_lang = data.get("target_lang", "es")

    if not session_id or session_id not in _store:
        return jsonify({"error": "Session not found. Please re-upload your file."}), 400

    entry      = _store[session_id]
    file_bytes = entry["bytes"]
    ext        = entry["ext"]
    filename   = entry["filename"]
    src        = entry.get("detected_code", "auto")

    bio = io.BytesIO(file_bytes)
    bio.name = filename

    one_fn  = _translate_one_fn(src, target_lang)
    many_fn = _translate_many_fn(src, target_lang)

    try:
        if ext == "pdf":
            result = translate_pdf_inplace(bio, one_fn, translate_many_fn=many_fn)
        elif ext == "docx":
            result = translate_docx_inplace(bio, one_fn, translate_many_fn=many_fn)
        elif ext in ("xlsx", "xls"):
            result = translate_xlsx_inplace(bio, one_fn, translate_many_fn=many_fn)
        else:
            # Plain text fallback
            text   = file_bytes.decode("utf-8", errors="ignore")
            xlated = translate_text(text, src, target_lang) or text
            result = io.BytesIO(xlated.encode("utf-8"))
    except Exception as exc:
        logging.error(f"Translation error: {exc}")
        return jsonify({"error": str(exc)}), 500

    if result is None:
        return jsonify({"error": "Translation failed — the file may be corrupt or unsupported."}), 500

    # result is a BytesIO; read its bytes for storage
    result.seek(0)
    result_bytes = result.read()

    preview = _extract_preview(result_bytes, ext)

    # Count pages
    pages = 1
    try:
        if ext == "pdf":
            import fitz
            doc   = fitz.open(stream=result_bytes, filetype="pdf")
            pages = len(doc)
            doc.close()
    except Exception:
        pass

    entry["translated_bytes"] = result_bytes
    entry["target_lang"]      = target_lang
    entry["target_name"]      = get_language_name(target_lang)
    entry["translated_ext"]   = ext
    entry["preview"]          = preview

    return jsonify({
        "preview":     preview,
        "target_lang": get_language_name(target_lang),
        "pages":       pages,
    })


@app.route("/api/download/<session_id>/<fmt>")
def download(session_id: str, fmt: str):
    if session_id not in _store:
        return jsonify({"error": "Session not found"}), 404

    entry = _store[session_id]
    translated_bytes: bytes | None = entry.get("translated_bytes")
    if translated_bytes is None:
        return jsonify({"error": "No translation available. Translate first."}), 400

    ext        = entry.get("translated_ext", "txt")
    ts         = datetime.now().strftime("%Y%m%d_%H%M")
    base_name  = f"sabio_translated_{ts}"
    fmt        = fmt.lower()

    # Return native format directly
    if fmt == ext:
        bio = io.BytesIO(translated_bytes)
        bio.seek(0)
        return send_file(bio, mimetype=_mime(fmt), as_attachment=True,
                         download_name=f"{base_name}.{fmt}")

    if fmt == "txt":
        text = _bytes_to_text(translated_bytes, ext)
        bio  = io.BytesIO(text.encode("utf-8"))
        return send_file(bio, mimetype=_mime("txt"), as_attachment=True,
                         download_name=f"{base_name}.txt")

    if fmt == "html":
        text = _bytes_to_text(translated_bytes, ext)
        lang = entry.get("target_name", "")
        html = (
            f"<!DOCTYPE html>\n<html lang=\"en\">\n<head>"
            f"<meta charset=\"utf-8\"><title>Sabio Translation — {lang}</title>"
            f"<style>body{{font-family:sans-serif;max-width:860px;margin:2rem auto;"
            f"padding:0 1rem;line-height:1.7}}</style></head>\n"
            f"<body><pre style=\"white-space:pre-wrap\">{text}</pre></body>\n</html>"
        )
        bio = io.BytesIO(html.encode("utf-8"))
        return send_file(bio, mimetype=_mime("html"), as_attachment=True,
                         download_name=f"{base_name}.html")

    if fmt == "docx" and ext != "docx":
        try:
            from utils.docx_reader import create_translated_docx
            text       = _bytes_to_text(translated_bytes, ext)
            docx_bytes = create_translated_docx(text)
            bio        = io.BytesIO(docx_bytes)
            return send_file(bio, mimetype=_mime("docx"), as_attachment=True,
                             download_name=f"{base_name}.docx")
        except Exception:
            pass

    # Fallback: return as plain text
    text = _bytes_to_text(translated_bytes, ext)
    bio  = io.BytesIO(text.encode("utf-8"))
    return send_file(bio, mimetype=_mime("txt"), as_attachment=True,
                     download_name=f"{base_name}.txt")


if __name__ == "__main__":
    app.run(debug=True, port=5000)
